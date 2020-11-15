#!/usr/bin/env python  
# encoding: utf-8  

""" 
@version: v1.0 
@author: xag 
@license: Apache Licence  
@contact: xinganguo@gmail.com 
@site: https://github.com/xingag 
@software: PyCharm 
@file: modify_.py 
@time: 2020/11/6 19:40 
@description：Word实战
"""

from docx import Document
from docxcompose.composer import Composer
import glob
from docx.shared import RGBColor
import os
from utils.docx_util import *
import codecs
from difflib import HtmlDiff
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne, ZeroOrMore, OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import ns


# 依赖
# pip3 install docxcompose

# 公众号：AirPython，欢迎关注！

class WordF(object):
    def __init__(self):
        self.path = './raw.docx'
        self.doc = Document(self.path)

    def modify(self):
        # 1、替换文字内容
        # 比如：将文档中所有的关键字“祖国”换成“父母”
        # self.replace_content("祖国", "父母")

        # 2、修改包含文字块的样式
        # 比如：把文字块包含“AirPython”的Run，进行加粗、标红处理
        # self.modify_run_style("AirPython")

        # 3、页眉页脚处理
        self.handle_header_and_footer()

        # 4、合并多个文件
        # self.compose_files(['new.docx', 'output.docx'], 'compose.docx')

        # 5、对比两个Word文档
        self.compare_files("./raw.docx", "./raw1.docx")

        # 6、新增数字索引
        doc = Document('./compose.docx')
        # 注意：要设置页眉页脚的对齐方式，必须设置到段落上（文字块不能添加对齐方式）
        doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 创建一个文字块样式，指定字体名称、大小、颜色
        style = create_style(document=doc, style_name="style", style_type=2, font_size=10,
                             font_color=[0x00, 0x00, 0x00], font_name="黑体")
        self.add_page_number(doc.sections[0].footer.paragraphs[0].add_run("", style))
        doc.save("./output.docx")
        print('添加页码索引成功！')

    def replace_content(self, old_content, new_content):
        """
        替换文档中所有内容
        :param old_content:旧的内容
        :param new_content:新的内容
        :return:
        """
        # 替换段落
        for paragraph in self.doc.paragraphs:
            if old_content in paragraph.text:
                # 替换内容后，重新设置进去
                paragraph.text = paragraph.text.replace(old_content, new_content)

        # 替换表格
        # document.tables[表格索引].rows[行索引].cells[单元格列索引].text = “新的数据”。
        tables = [table for table in self.doc.tables]
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_content in cell.text:
                        # 重新设置单元格内容
                        cell.text = cell.text.replace(old_content, new_content)

        # 保存到一个新的文件中
        self.doc.save('./new.docx')

    def modify_run_style(self, keyword):
        """
        修改文字块样式
        :return:
        """
        # 获取当前目录下所有的Word文件
        for file in glob.glob('./*.docx'):
            docx = Document(file)

            # 关键字的文字块或单元格标红，并加粗
            # 1、修改段落中包含关键字的文件块的样式
            for paragraph in docx.paragraphs:
                for run in paragraph.runs:
                    if keyword in run.text:
                        # 修改颜色为红色，并加粗显示
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)

            # 2、修改表格中，包含关键字的某个单元格的样式
            # 注意：只有add_run的元素才有font属性（新增的表格）
            tables = [table for table in docx.tables]
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        if keyword in cell.text:
                            # 原内容
                            content_raw = cell.text
                            # 清空单元格数据
                            cell.text = ""
                            # 追加数据进去，并设置样式
                            run = cell.paragraphs[0].add_run(content_raw)
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.bold = True
            docx.save('./output/' + os.path.basename(file))

    def handle_header_and_footer(self):
        """
        处理页眉页脚
        包含：新增页眉、页脚
        :return:
        """
        # 1、获取待处理页眉、页脚的章节
        header = self.doc.sections[0].header
        footer = self.doc.sections[0].footer

        # True if this section displays a distinct first-page header and footer
        # True：页眉页脚不同于首页，每个页面章节的页眉页脚单独设置
        # False：每个页面的页眉页脚相同
        # self.doc.sections[0].different_first_page_header_footer = True

        # 2、新增页眉
        # 2.1 普通的页眉、页脚
        # add_norm_header_and_footer(header, footer, "我是一个页眉", "我是一个页脚")

        # 注意：要设置页眉页脚的对齐方式，必须设置到段落上，使用document.styles.add_style无效
        # header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 2.2 自带样式的页眉、页脚
        # 创建一个样式
        # style_paragraph = create_style(document=self.doc, style_name="style5", style_type=2, font_size=30,
        #                                font_color=[0xff, 0x00, 0x00])
        # add_custom_style_header_and_footer(header, footer, "我是页眉2", "我是页脚2", style_paragraph)

        # 3、删除页眉、页脚
        remove_all_header_and_footer(self.doc)

        # 4、重新保存
        self.doc.save(self.path)

    def compose_files(self, files, output_file_path):
        """
        合并多个word文件到一个文件中
        :param files:
        :return:
        """
        composer = Composer(Document())
        for file in files:
            composer.append(Document(file))

        # 保存到新的文件中
        composer.save(output_file_path)

    def compare_files(self, param, param1):
        """
        对比两个word文档，并生成差异性报告
        :param param:
        :param param1:
        :return:
        """
        file1 = Document(param)
        file2 = Document(param1)

        # 分别获取段落内容
        content1 = ''
        content2 = ''
        for paragraph in file1.paragraphs:
            if "" == paragraph.text.strip():
                continue
            content1 += paragraph.text + '\n'

        for paragraph in file2.paragraphs:
            if "" == paragraph.text.strip():
                continue
            content2 += paragraph.text + '\n'

        # 如果参数 keepends 为 False，不包含换行符，如果为 True，则保留换行符。
        print("第二个文档数据如下：\n", content1.splitlines(keepends=False))
        print("第一个文档数据如下：\n", content1.splitlines(keepends=False))

        diff_html = HtmlDiff(wrapcolumn=100).make_file(content1.split("\n"), content2.split("\n"))
        with codecs.open('./diff_result.html', 'w', encoding='utf-8') as f:
            f.write(diff_html)

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        """
        添加页面索引
        :param run:
        :return:
        """
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        # run._r：class 'docx.oxml.text.run.CT_R'>
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


if __name__ == '__main__':
    wordF = WordF()
    wordF.modify()
