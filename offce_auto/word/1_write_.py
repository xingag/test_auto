#!/usr/bin/env python  
# encoding: utf-8  

""" 
@version: v1.0 
@author: xag 
@license: Apache Licence  
@contact: xinganguo@gmail.com 
@site: https://github.com/xingag 
@software: PyCharm 
@file: main.py 
@time: 2020/11/5 19:24 
@description：Python写入数据到Word中
"""

from docx import Document
from font import FontRun
from utils.docx_util import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 依赖
# pip3 install python-docx


class WordF(object):
    def __init__(self):
        pass

    def write(self):
        """
        写入数据到Word文件中
        :return:
        """
        # 1、新建一个空白文档
        doc = Document()

        # 2、新增内容
        # 2.1 标题
        # 分别写入一个一级标题，一个二级标题，一个三级标题
        doc.add_heading('一级标题', 0)
        doc.add_heading('二级标题', 1)
        doc.add_heading('三级标题', 2)

        # 2.2 段落
        # 2.2.1 新增普通段落
        doc.add_paragraph("我是一个普通段落。")

        # 2.2.2.1 新增一个带字体样式的段落（方式一）
        paragraph = doc.add_paragraph()
        FontRun(paragraph, "我是一个自带样式的段落（方式一）！！！").set_font_size(40).set_bold(True).set_underline(True).set_italic(
            True).set_font_color([0xff, 0x00, 0x00]).set_font_name('Times New Roman')

        # 2.2.2.2 新增一个带字体样式的段落（方式二）
        # 1/段落样式
        style_paragraph = create_style(document=doc, style_name="style2", style_type=1, font_size=30,
                                       font_color=[0xff, 0x00, 0x00])
        # 2/字符样式
        style_string = create_style(document=doc, style_name="style3", style_type=2, font_size=15,
                                    font_color=[0x00, 0xff, 0x00])
        # 3/表格样式
        # 对齐方式为：居中
        style_table = create_style(document=doc, style_name="style4", style_type=3, font_size=25,
                                   font_color=[0x00, 0x00, 0xff], align=WD_PARAGRAPH_ALIGNMENT.CENTER)

        current_paragraph = doc.add_paragraph("我是一个自带样式的段落（方式二）！！！", style_paragraph)
        # 字符样式

        current_paragraph.add_run("【段落2中的部分字符】", style_string)

        # 2.2.3 新增一个引用段落
        doc.add_paragraph('--我是一个引用段落--', style='Intense Quote')

        # 2.3 列表
        # 2.3.1 无序列表
        add_list(doc, ["无序-Item1", "无序-Item2", "无序-Item3"], False)

        # 2.3.2 有序列表
        add_list(doc, ["有序-Item1", "有序-Item2", "有序-Item3"], True)

        # 在word中，图片和表格使用频率也很高
        # 2.4 图片
        # 2.4.1 插入本地图片
        add_local_image(doc, './1.png', width=2)

        # 2.4.2 插入网络图片
        add_network_image(doc, url, width=3)

        # 2.5 表格
        head_datas = ["姓名", "年龄", "地区"]
        datas = (
            ('张三', 18, '深圳'),
            ('李四', 28, '北京'),
            ('王五', 33, '上海'),
            ('孙六', 42, '广州')
        )

        # 新增一个表格，并指定样式
        # add_table(doc, head_datas, datas, style_table)
        add_table(doc, head_datas, datas)

        # 保存文件
        doc.save('output.docx')

if __name__ == '__main__':
    wordF = WordF()
    wordF.write()
