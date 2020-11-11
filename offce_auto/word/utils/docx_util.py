#!/usr/bin/env python  
# encoding: utf-8  

""" 
@version: v1.0 
@author: xag 
@license: Apache Licence  
@contact: xinganguo@gmail.com 
@site: https://github.com/xingag 
@software: PyCharm 
@file: docx_util.py 
@time: 2020/11/5 23:24 
@description：工具类
"""

from docx.shared import Inches
from urllib.request import urlopen
from io import BytesIO
import ssl
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re


# 公众号：AirPython

def get_image_data_from_network(url):
    """
    获取网络图片字节流
    :param url: 图片地址
    :return:
    """
    ssl._create_default_https_context = ssl._create_unverified_context
    # 获取网络图片的字节流
    image_data = BytesIO(urlopen(url).read())
    return image_data


def add_list(document, data, isorder):
    """
    将列表数据添加到无序列表/有序列表中
    :param document: 文档对象
    :param data: 列表数据
    :param isorder: 是否有序列表
    :return:
    """
    # 无序列表
    if not isorder:
        for item in data:
            document.add_paragraph(item, style='List Bullet')
    else:
        # 有序列表
        for item in data:
            document.add_paragraph(item, style='List Number')


def add_local_image(doc, image_path, width=None, height=None):
    """
    增加本地图片到Word文档中
    :param doc:
    :param image_path:
    :param width:
    :param height:
    :return:
    """
    doc.add_picture(image_path, width=None if width is None else Inches(width),
                    height=None if height is None else Inches(height))


def add_network_image(doc, image_url, width=None, height=None):
    """
    增加本地图片到Word文档中
    :param doc:
    :param image_url:
    :param width:
    :param height:
    :return:
    """
    # 获取图片流
    image_data = get_image_data_from_network(image_url)
    doc.add_picture(image_data, width=None if width is None else Inches(width),
                    height=None if height is None else Inches(height))


def add_table(doc, head_datas, datas, style=None):
    """
    新增一个表格
    :param doc:
    :param head_datas: 表头
    :param datas: 数据
    :param style:
    :return:
    """
    # 新增一个表格
    # 默认样式为：Table Grid
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))

    # 第一行所有单元格对象列表
    head_cells = table.rows[0].cells

    # 写入数据到表头中
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item

    # 遍历数据并写入数据
    for data in datas:
        # 单独添加一行或者列：add_row、add_column
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])


def create_style(document, style_name, style_type, font_size=-1, font_color=None, font_name=None, align=None):
    """
    创建一个样式
    :param align:
    :param document:
    :param style_name: 样式名称
    :param style_type: 样式类型，1：段落样式, 2：字符样式, 3：表格样式
    :param font_name:
    :param font_color:
    :param font_size:
    :return:
    """
    if font_color is None:
        font_color = []

    # 注意：必须要判断样式是否存在，否则重新添加会报错
    style_names = [style.name for style in document.styles]
    if style_name in style_names:
        # print('样式已经存在，不需要重新添加！')
        return

    font_style = document.styles.add_style(style_name, style_type)

    # 字体大小
    if font_size != -1:
        font_style.font.size = Pt(font_size)

    # 字体颜色
    # 比如：[0xff,0x00,0x00]
    if font_color and len(font_color) == 3:
        font_style.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])

    # 对齐方式
    # 注意：段落、表格才有对齐方式
    if style_type != 2 and align:
        font_style.paragraph_format.alignment = align
        # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 中文字体名称
    if font_name:
        font_style.font.name = font_name
        font_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    return font_style


def get_page_margin(section):
    """
    获取某个页面的页边距(EMU)
    :param section:
    :return:
    """
    # 分别对应：左边距、上边距、右边距、下边距
    left, top, right, bottom = section.left_margin, section.top_margin, section.right_margin, section.bottom_margin
    return left, top, right, bottom

def get_header_footer_distance(section):
    """
    获取页眉、页脚边距
    :param section:
    :return:
    """
    # 分别对应页眉边距、页脚边距
    header_distance, footer_distance = section.header_distance, section.footer_distance
    return header_distance, footer_distance

def get_page_size(section):
    """
    获取页面宽度、高度
    :param section:
    :return:
    """
    # 分别对应页面宽度、高度
    page_width, page_height = section.page_width, section.page_height
    return page_width, page_height


def get_page_orientation(section):
    """
    获取页面方向
    :param section:
    :return:
    """
    return section.orientation


def get_runs(paragraph):
    """
    获取段落下所有的文字块信息，包含：数目、内容列表
    :param paragraph:
    :return:
    """
    # 段落对象包含的文字块Run
    runs = paragraph.runs

    # 数量
    runs_length = len(runs)

    # 文字块内容
    runs_contents = [run.text for run in runs]

    return runs, runs_length, runs_contents

def get_table_cell_content(table):
    """
    读取表格中所有单元格是内容
    :param table:
    :return:
    """
    # 所有单元格
    cells = table._cells
    cell_size = len(cells)

    # 所有单元格的内容
    content = [cell.text for cell in cells]

def get_table_size(table):
    """
    获取表格的行数量、列数量
    :param table:
    :return:
    """
    # 几行、几列
    row_length, column_length = len(table.rows), len(table.columns)
    return row_length, column_length

def get_table_row_datas(table):
    """
    获取表格中行数据
    :param table:
    :return:
    """
    rows = table.rows
    datas = []

    # 每一行获取单元格的数据组成列表，加入到结果列表中
    for row in rows:
        datas.append([cell.text for cell in row.cells])
    return datas


def get_table_column_datas(table):
    """
    获取表格中列数据
    :param table:
    :return:
    """
    columns = table.columns
    datas = []

    # 每一列获取单元格的数据组成列表，加入到结果列表中
    for column in columns:
        datas.append([cell.text for cell in column.cells])
    return datas


def get_word_pics(doc, word_path, output_path):
    """
    提取word文档内的图片
    :param word_path:源文件名称
    :param output_path: 结果目录
    :return:
    """
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            # 图片保存目录
            if not os.path.exists(output_path):
                os.makedirs(output_path)
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]

            # 新的名称
            newname = word_name.split('\\')[-1] if os.sep in word_name else word_name.split('/')[-1]
            img_name = f'{newname}_{img_name}'

            # 写入到文件中
            with open(f'{output_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)





