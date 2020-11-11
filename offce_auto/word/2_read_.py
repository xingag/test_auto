#!/usr/bin/env python  
# encoding: utf-8  

""" 
@version: v1.0 
@author: xag 
@license: Apache Licence  
@contact: xinganguo@gmail.com 
@site: https://github.com/xingag 
@software: PyCharm 
@file: read_.py 
@time: 2020/11/6 18:39 
@description：读取word文件
"""

from docx import Document
from docx.enum.section import WD_ORIENT

from utils.docx_util import *


# 公众号：AirPython


class WordF(object):

    def __init__(self):
        # 源文件目录
        self.word_path = './output.docx'

        # 打开文档，构建一个文档对象
        self.doc = Document(self.word_path)

    def read(self):
        # 1、读取文档基本信息
        # self.read_basic_msg()

        # 2、段落信息
        # self.read_paragraph_msg()

        # 3、获取文字块信息
        # self.read_run_msg()

        # 4、获取word文档中的表格信息
        # self.read_table_msg()

        # 5、下载word中所有图片
        # 解压后的目录：word/media
        # get_word_pics(self.doc, self.word_path, "./output/")

        # 6、读取页面、页脚信息
        self.read_header_and_footer()

    def read_basic_msg(self):
        """
        读取Word文档基本信息
        :return:
        """
        # 1、获取章节信息
        # 注意：章节可以设置本页的大小、页眉、页脚
        msg_sections = self.doc.sections
        print("章节列表：", msg_sections)
        # 章节数目
        print('章节数目:', len(msg_sections))

        # 2、页边距信息
        first_section = msg_sections[0]
        left, top, right, bottom = get_page_margin(first_section)
        print('左边距:', left, ",上边距：", top, ",右边距：", right, "，下边距：", bottom)

        # 3、页眉页脚边距
        header_distance, footer_distance = get_header_footer_distance(first_section)
        print('页眉边距:', header_distance, "，页脚边距:", footer_distance)

        # 4、页面宽度、高度
        page_width, page_height = get_page_size(first_section)
        print('页面宽度:', page_width, "，页面高度:", page_height)

        # 5、页面方向
        # 类型：class 'docx.enum.base.EnumValue
        # 包含：PORTRAIT (0)、LANDSCAPE (1)
        page_orientation = get_page_orientation(first_section)
        print("页面方向:", page_orientation)

        # 设置页面方向（横向、竖向）
        # 设置为横向
        first_section.orientation = WD_ORIENT.LANDSCAPE
        # 设置为竖向
        # first_section.orientation = WD_ORIENT.PORTRAIT
        self.doc.save(self.word_path)

    def read_header_and_footer(self):
        """
        读取页眉页脚信息
        :return:
        """
        # 获取某一个章节
        first_section = self.doc.sections[0]

        # 注意：页眉、页脚都有可能包含多个段落
        # 页眉所有的段落
        header_content = " ".join([paragraph.text for paragraph in first_section.header.paragraphs])
        print("页眉内容:", header_content)

        # 页脚
        footer_content = " ".join([paragraph.text for paragraph in first_section.footer.paragraphs])
        print("页脚内容:", footer_content)

    def read_paragraph_msg(self):
        """
        获取段落信息
        :return:
        """
        # 获取文档对象中所有的段落，默认不包含：页眉、页脚、表格中的段落
        paragraphs = self.doc.paragraphs

        # 0、读取所有段落数据
        contents = [paragraph.text for paragraph in self.doc.paragraphs]
        print(contents)

        # 1、段落数目
        paragraphs_length = len(paragraphs)
        print('文档中一共包含：{}个段落'.format(paragraphs_length))

        # 2、获取某一个段落的格式信息
        paragraph_someone = paragraphs[0]

        # 2.1 段落内容
        content = paragraph_someone.text
        print('段落内容:', content)

        # 2.2 段落格式
        paragraph_format = paragraph_someone.paragraph_format

        # 2.2.1 对齐方式
        # <class 'docx.enum.base.EnumValue'>
        alignment = paragraph_format.alignment
        print('段落对齐方式：', alignment)

        # 2.2.2 左、右缩进
        left_indent, right_indent = paragraph_format.left_indent, paragraph_format.right_indent
        print('段落左缩进:', left_indent, "，右缩进:", right_indent)

        # 2.2.3 首行缩进
        first_line_indent = paragraph_format.first_line_indent
        print('段落首行缩进:', first_line_indent)

        # 2.2.4 行间距
        line_spacing = paragraph_format.line_spacing
        print('段落行间距:', line_spacing)

        # 2.2.5 段落前后间距
        space_before, space_after = paragraph_format.space_before, paragraph_format.space_after
        print('段落前、后间距分别为:', space_before, ',', space_after)

    def read_run_msg(self):
        """
        获取文字块基本信息（即：Run对象），包含：内容、文字格式信息
        注意：一个段落包含多个文字块
        :return:
        """
        # 某一个段落对象
        paragraph_someone = self.doc.paragraphs[0]

        # 1、段落下的文字块基本信息（包含：换行符等）
        runs, runs_length, runs_contents = get_runs(paragraph_someone)
        print('文字块数目：', runs_length, '\n内容分别是：', runs_contents)

        # 2、文字块格式信息
        # 包含：字体名称、大小、颜色、是否加粗等
        # 某一个文字块的字体属性
        run_someone_font = runs[0].font

        # 字体名称
        font_name = run_someone_font.name
        print('字体名称:', font_name)

        # 字体颜色(RGB)
        # <class 'docx.shared.RGBColor'>
        font_color = run_someone_font.color.rgb
        print('字体颜色:', font_color)
        print(type(font_color))

        # 字体大小
        font_size = run_someone_font.size
        print('字体大小:', font_size)

        # 是否加粗
        # True:加粗；None/False：没有加粗
        font_bold = run_someone_font.bold
        print('是否加粗：', font_bold)

        # 是否斜体
        # True:协议；None/False：不是斜体
        font_italic = run_someone_font.italic
        print('是否斜体:', font_italic)

        # 带下划线
        # True：带有下滑线；None/False:字体没有带下滑线
        font_underline = run_someone_font.underline
        print('带有下滑线:', font_underline)

        # 删除线/双删除线
        # True：带有删除线；None/False:字体没有带删除线
        font_strike = run_someone_font.strike
        font_double_strike = run_someone_font.double_strike
        print('带有删除线:', font_strike, "\n带有双删除线:", font_double_strike)

    def read_table_msg(self):
        """
        获取表格信息
        :return:
        """
        # 文档中所有的表格对象
        tables = self.doc.tables

        # 1、表格数量
        table_num = len(tables)
        print('文档中包含的表格数量:', table_num)

        # 2、读取所有表格数据
        # 所有表格对象
        print('内容分别是：')
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end='  ')
                print()
            print('\n')

        # 3、表格样式名称
        # Table Grid
        table_someone = tables[0]
        style = table_someone.style.name
        print("表格样式:", style)

        # 4、获取某一个表格中所有单元格的内容
        content = get_table_cell_content(table_someone)
        print('表格中所有单元格数据如下:', content)

        # 5、表格行、列数量
        row_length, column_length = get_table_size(table_someone)
        print('表格包含：%d行%d列' % (row_length, column_length))

        # 6、表格行数据
        rows_datas = get_table_row_datas(table_someone)
        print('按行组成的数据如下:\n', rows_datas)

        # 7、表格列数据
        column_datas = get_table_column_datas(table_someone)
        print('按列组成的数据如下:\n', column_datas)


if __name__ == '__main__':
    wordF = WordF()
    wordF.read()
